"""Force RTL on ALL table cells to prevent BIDI flipping of English at sentence end."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("template_essic.docx")

def set_rtl_paragraph(p):
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.insert(0, bidi)

# Force RTL on every paragraph in every table cell
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_rtl_paragraph(p)

doc.save("template_essic.docx")
print("All table cells forced RTL.")
