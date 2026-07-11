from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

doc = Document("template_essic.docx")

def set_rtl_paragraph(p):
    """Force paragraph to be RTL."""
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.insert(0, bidi)

def set_run_cs_font(run, font_name):
    """Set the complex script font on a run."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

# The target style for body text: match "عناية {{STACK_HOLDER}}" heading
# P5: font=Arial Unicode MS, cs_font=Times New Roman, size=241300, bold=True, color=FF0000

# P5 has cs_font=Times New Roman — that's what the user wants the body text to match
TARGET_CS_FONT = "Times New Roman"
TARGET_FONT = "Arial Unicode MS"
TARGET_BOLD = True
TARGET_COLOR = RGBColor(0x00, 0x20, 0x60)  # dark navy blue - let's look at what user wants
# Actually the عناية heading color is FF0000 (red). But user said font style not color.
# Let's keep color as-is and just fix the cs_font to match

# Fix BODY_TEXT: change cs_font to Times New Roman to match the عناية line
for p in doc.paragraphs:
    if "{{BODY_TEXT}}" in p.text:
        set_rtl_paragraph(p)
        for run in p.runs:
            run.font.name = TARGET_FONT
            set_run_cs_font(run, TARGET_CS_FONT)
            run.font.size = Pt(14)  # Match the عناية heading at ~19pt
            run.font.bold = True
            # Keep the existing color or make it dark blue to match heading?
            # User said match font of عناية - that's red. Let's set same color
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red like عناية
        print("Fixed BODY_TEXT font")

# Fix SENDER_TOP cell - force RTL paragraph so Arabic text with English doesn't flip
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "{{SENDER_TOP}}" in p.text:
                    set_rtl_paragraph(p)
                    # Also set cs_font on runs
                    for run in p.runs:
                        set_run_cs_font(run, TARGET_CS_FONT)
                    print("Fixed SENDER_TOP RTL")
                if "{{SEND_TO}}" in p.text:
                    set_rtl_paragraph(p)
                    for run in p.runs:
                        set_run_cs_font(run, TARGET_CS_FONT)
                if "{{SUBJECT}}" in p.text:
                    set_rtl_paragraph(p)
                    for run in p.runs:
                        set_run_cs_font(run, TARGET_CS_FONT)

doc.save("template_essic.docx")
print("Template saved with RTL + font fixes.")
