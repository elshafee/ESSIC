from docx import Document
doc = Document("template_essic.docx")
print("Paragraphs containing {{SENDER}}:")
for i, p in enumerate(doc.paragraphs):
    if "{{SENDER}}" in p.text:
        print(f"P{i}: {p.text}")
print("Tables containing {{SENDER}}:")
for i, t in enumerate(doc.tables):
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if "{{SENDER}}" in p.text:
                    print(f"Table {i}: {p.text}")
