from docx import Document
from docx.oxml.ns import qn

doc = Document("template_essic.docx")

print("=== ALL PARAGRAPHS FONT INFO ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        for r in p.runs:
            if r.text.strip():
                rPr = r._element.find(qn('w:rPr'))
                cs_font = None
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        cs_font = rFonts.get(qn('w:cs'))
                color = r.font.color.rgb if (r.font.color and r.font.color.type) else None
                print(f"P{i} [{p.text[:30]}]: font={r.font.name}, cs_font={cs_font}, size={r.font.size}, bold={r.font.bold}, color={color}")
                break

print("\n=== TABLE PARAGRAPHS ===")
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                if p.text.strip():
                    for run in p.runs:
                        if run.text.strip():
                            rPr = run._element.find(qn('w:rPr'))
                            cs_font = None
                            if rPr is not None:
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    cs_font = rFonts.get(qn('w:cs'))
                            print(f"T{t_idx}[r{r_idx}c{c_idx}] [{p.text[:30]}]: cs_font={cs_font}, font={run.font.name}")
                            break
