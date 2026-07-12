from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

doc = Document("template_essic.docx")

# First, find the style of the manager text
manager_color = None
manager_bold = None
manager_size = None
manager_font = None

for p in doc.paragraphs:
    if "مدير مركز الخدمات الإلكترونية" in p.text and "(ESSIC)" in p.text:
        # Check if this is the RED one at the bottom
        for r in p.runs:
            if r.text.strip():
                if r.font.color and r.font.color.rgb:
                    manager_color = r.font.color.rgb
                else:
                    # sometimes color is not explicitly on run but paragraph
                    pass
                manager_bold = r.font.bold
                manager_size = r.font.size
                manager_font = r.font.name
        break

print(f"Manager style - Color: {manager_color}, Bold: {manager_bold}, Size: {manager_size}, Font: {manager_font}")

# Now apply this style to {{BODY_TEXT}}
for p in doc.paragraphs:
    if "{{BODY_TEXT}}" in p.text:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Arial Unicode MS'
            # Also set the complex script font
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'), 'Arial Unicode MS')
            # Let's set the size to something readable but prominent (size 14 pt is standard for bold headers)
            # The user asked to match the manager exactly. Let's see if we can just copy it:
            if manager_size:
                r.font.size = manager_size
            else:
                r.font.size = Pt(14)
            
            if manager_color:
                r.font.color.rgb = manager_color
            else:
                # If the top one was found, it might be black. Let's force red if they meant the signature?
                # Actually, let's just make it Bold and Black if it's the top one, or Red if it's the bottom.
                pass
        print("Updated BODY_TEXT style")

doc.save("template_essic.docx")
print("Template saved.")
