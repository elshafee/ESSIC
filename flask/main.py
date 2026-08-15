import base64
import io
import os
import sys
from datetime import datetime, timedelta
from functools import wraps

import pyotp
import qrcode
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, abort, session, \
    jsonify
from sqlalchemy import true
from werkzeug.utils import secure_filename

# Allow imports from project root
sys.path.insert(0, os.path.dirname(__file__))

from models import db, Document, User, Contact
from services.numbering import get_next_serial, build_full_code
from services.word_editor import replace_placeholder

# Load environment variables
load_dotenv()

# ─── App Configuration ────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
import tempfile
if os.environ.get("VERCEL") == "1":
    DATA_DIR = os.environ.get("DATA_DIR", tempfile.gettempdir())
else:
    DATA_DIR = os.environ.get("DATA_DIR", BASE_DIR)
UPLOAD_FOLDER = os.path.join(DATA_DIR, "uploads")
GENERATED_FOLDER = os.path.join(DATA_DIR, "generated")
ALLOWED_EXT = {"docx"}

app = Flask(__name__)
app.secret_key = "essic-doc-numbering-secret-2026"
app.permanent_session_lifetime = timedelta(days=30)
supabase_url = os.environ.get("SUPABASE_DB_URL")
if not supabase_url:
    raise RuntimeError("SUPABASE_DB_URL environment variable is required and not set.")

if supabase_url.startswith("postgres://"):
    supabase_url = supabase_url.replace("postgres://", "postgresql+psycopg://", 1)
elif supabase_url.startswith("postgresql://"):
    supabase_url = supabase_url.replace("postgresql://", "postgresql+psycopg://", 1)
app.config["SQLALCHEMY_DATABASE_URI"] = supabase_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["GENERATED_FOLDER"] = GENERATED_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20 MB limit

# Ensure folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

_admin_raw = os.environ.get("ADMIN_EMAILS", "mkhassan@horus.edu.eg,aelshafee@horus.edu.eg")
ADMIN_EMAILS = {e.strip().lower() for e in _admin_raw.split(",") if e.strip()}

# Allowed users — comma-separated emails in .env
_allowed_raw = os.environ.get("ALLOWED_EMAILS", "mkhassan@horus.edu.eg,aelshafee@horus.edu.eg")
ALLOWED_EMAILS = {e.strip().lower() for e in _allowed_raw.split(",") if e.strip()}


def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    return decorated_function


@app.context_processor
def inject_globals():
    is_admin = False
    if 'user' in session:
        user_email = session['user'].get('email', '')
        if user_email in ADMIN_EMAILS:
            is_admin = True
        else:
            user = User.query.filter_by(email=user_email).first()
            if user and user.role == 'admin':
                is_admin = True
    
    from services.onedrive import get_word_online_url
    return dict(is_admin=is_admin, get_word_online_url=get_word_online_url)


def is_current_user_admin():
    user_email = session.get('user', {}).get('email', '')
    if not user_email: return False
    if user_email in ADMIN_EMAILS:
        return True
    user = User.query.filter_by(email=user_email).first()
    return user and user.role == 'admin'


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not is_current_user_admin():
            flash("You do not have permission to access this page.", "danger")
            return redirect(url_for('documents'))
        return f(*args, **kwargs)

    return decorated_function


db.init_app(app)

with app.app_context():
    db.create_all()


# ─── Helpers ─────────────────────────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user' in session:
        return redirect(url_for('index'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()

        if not email:
            flash("Please enter your email.", "warning")
            return redirect(url_for('login'))

        # Check if user exists in DB
        user = User.query.filter_by(email=email).first()

        if not user and email not in ALLOWED_EMAILS:
            flash("This email address is not authorized to access this system. Please ask an admin to add you.",
                  "danger")
            return redirect(url_for('login'))

        if not user:
            # First time logging in: Create user and generate TOTP secret
            name = email.split('@')[0]
            totp_secret = pyotp.random_base32()
            is_admin = email in ADMIN_EMAILS
            user = User(email=email, name=name, totp_secret=totp_secret, is_setup=False,
                        role='admin' if is_admin else 'user')
            db.session.add(user)
            db.session.commit()

        session['pending_email'] = email

        if not user.is_setup:
            return redirect(url_for('setup'))
        else:
            return redirect(url_for('verify'))

    return render_template('login.html')


@app.route('/setup', methods=['GET', 'POST'])
def setup():
    if 'user' in session:
        return redirect(url_for('index'))

    email = session.get('pending_email')
    if not email:
        return redirect(url_for('login'))

    user = User.query.filter_by(email=email).first()
    if not user or user.is_setup:
        return redirect(url_for('login'))

    if request.method == 'POST':
        code = request.form.get('code', '').strip()
        totp = pyotp.TOTP(user.totp_secret)
        if totp.verify(code):
            user.is_setup = True
            db.session.commit()

            session.pop('pending_email', None)
            session['user'] = {'name': user.name, 'email': user.email}
            session.permanent = True
            return redirect(url_for('index'))
        else:
            flash("Invalid code. Please try again.", "danger")

    # Generate QR code URL
    totp_uri = pyotp.totp.TOTP(user.totp_secret).provisioning_uri(name=email, issuer_name="ESSIC Document System")

    # Generate QR code image as base64
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(totp_uri)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    qr_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

    return render_template('setup.html', qr_b64=qr_b64, secret=user.totp_secret, email=email)


@app.route('/verify', methods=['GET', 'POST'])
def verify():
    if 'user' in session:
        return redirect(url_for('index'))

    email = session.get('pending_email')
    if not email:
        return redirect(url_for('login'))

    user = User.query.filter_by(email=email).first()
    if not user or not user.is_setup:
        return redirect(url_for('login'))

    if request.method == 'POST':
        code = request.form.get('code', '').strip()
        totp = pyotp.TOTP(user.totp_secret)
        if totp.verify(code):
            session.pop('pending_email', None)
            session['user'] = {'name': user.name, 'email': user.email}
            session.permanent = True
            return redirect(url_for('index'))
        else:
            flash("Invalid authenticator code. Please try again.", "danger")

    return render_template('verify.html', email=email)


@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))


@app.route("/", methods=["GET", "POST"])
@login_required
def index():
    """Upload page — user uploads a .docx template and gets a numbered document, or generates directly from default template."""
    if request.method == "POST":
        action_type = request.form.get("action_type", "upload")

        if action_type == "upload":
            file = request.files.get("docx_file")

            if not file or file.filename == "":
                flash("Please select a Word document (.docx) to upload.", "warning")
                return redirect(url_for("index"))

            if not allowed_file(file.filename):
                flash("Only .docx files are accepted.", "danger")
                return redirect(url_for("index"))

            file_title = request.form.get("file_title", "").strip()

            # Save uploaded file
            original_name = secure_filename(file.filename)
            upload_path = os.path.join(UPLOAD_FOLDER, original_name)
            file.save(upload_path)

        elif action_type == "direct_template":
            file_title = request.form.get("file_title", "Direct Template Report").strip()
            upload_path = os.path.join(BASE_DIR, "template_essic.docx")
            original_name = "template_essic.docx"

            if not os.path.exists(upload_path):
                flash("Default template (template_essic.docx) not found.", "danger")
                return redirect(url_for("index"))
        elif action_type == "reserve":
            file_title = request.form.get("file_title", "Reserved Serial").strip()
            upload_path = os.path.join(BASE_DIR, "template_essic.docx")
            original_name = "template_essic.docx"

            if not os.path.exists(upload_path):
                flash("Default template (template_essic.docx) not found.", "danger")
                return redirect(url_for("index"))
        else:
            flash("Invalid action.", "danger")
            return redirect(url_for("index"))

        # Generate document number
        now = datetime.now()
        month, year = now.month, now.year
        serial = get_next_serial(month, year)
        full_code = build_full_code(serial, month, year)

        # Process Word document
        generated_name = f"{full_code.replace(' ', '_')}.docx"

        generated_path = os.path.join(GENERATED_FOLDER, generated_name)

        try:
            if action_type == "upload":
                replace_placeholder(upload_path, generated_path, full_code)
            elif action_type == "reserve":
                from services.word_editor import replace_placeholders
                replace_placeholders(upload_path, generated_path, {
                    "{{CODE_NUMBER}}": full_code,
                    "{{SEND_TO}}": " ",
                    "{{SUBJECT}}": " ",
                    "{{STACK_HOLDER}}": " ",
                    "{{POSITION}}": " ",
                    "{{BODY_TEXT}}": " ",
                    "{{SENDER}}": " ",
                    "{{SENDER_TOP}}": " ",
                    "{{ESSIC}}": " "
                })
            elif action_type == "direct_template":
                from services.word_editor import replace_placeholders
                from services.ai_generator import normalize_recipient

                recipient = request.form.get("recipient", "")
                subject = request.form.get("subject", "")
                holder_name = request.form.get("holder_name", "")
                position = request.form.get("position", "")
                full_body = request.form.get("full_body", "")
                sender = request.form.get("sender", "")
                sender_position = request.form.get("sender_position", "")
                is_internal = request.form.get("is_internal") == "1"
                manager_is_sender = request.form.get("manager_is_sender") == "1"

                tables_data_str = request.form.get("tables_data", "[]")
                import json
                try:
                    tables_data = json.loads(tables_data_str)
                except:
                    tables_data = []

                image_paths = []
                images = request.files.getlist("images")
                if images:
                    import uuid
                    for img in images:
                        if img.filename:
                            ext = img.filename.split('.')[-1]
                            temp_path = os.path.join(UPLOAD_FOLDER, f"temp_{uuid.uuid4().hex}.{ext}")
                            img.save(temp_path)
                            image_paths.append(temp_path)

                import re
                def remove_office_code(text):
                    if not text: return ""
                    text = re.sub(r'\(\s*ESSIC\s*\)', '', text, flags=re.IGNORECASE)
                    text = re.sub(r'ESSIC', '', text, flags=re.IGNORECASE)
                    return text.strip()

                default_sender = 'مدير مركز الخدمات الإلكترونية والإبداع العلمي (ESSIC)'
                
                # Unconditionally remove ESSIC text from all text fields
                recipient = remove_office_code(recipient)
                subject = remove_office_code(subject)
                holder_name = remove_office_code(holder_name)
                position = remove_office_code(position)
                full_body = remove_office_code(full_body)
                sender = remove_office_code(sender)
                sender_position = remove_office_code(sender_position)
                default_sender = remove_office_code(default_sender)

                # Determine sender value for template
                sender_value = ""
                sender_name_used = sender if sender else default_sender
                if sender_position:
                    sender_top = f"\u202A{sender_name_used} / {sender_position}\u202C"
                else:
                    sender_top = f"\u202A{sender_name_used}\u202C"

                # If external OR manager is sender -> sender_value stays empty (manager signature kept)
                # Only if internal AND NOT manager -> employee signature used
                if is_internal and not manager_is_sender and sender:
                    sender_value = f"\u202Bمقدمه لسيادتكم\n{sender}\u202C"

                remove_manager = (is_internal and not manager_is_sender)
                replace_placeholders(upload_path, generated_path, {
                    "{{CODE_NUMBER}}": full_code,
                    "{{SEND_TO}}": ("\u202A" + normalize_recipient(recipient) + "\u202C") if recipient else "",
                    "{{SUBJECT}}": "\u202A" + subject + "\u202C" if subject else " ",
                    "{{STACK_HOLDER}}": "\u202B" + holder_name + "\u202C",
                    "{{POSITION}}": "\u202B" + position + "\u202C",
                    "{{BODY_TEXT}}": "\u202B" + full_body + "\u202C",
                    "{{SENDER}}": sender_value,
                    "{{SENDER_TOP}}": sender_top,
                    "{{ESSIC}}": "",
                }, remove_manager_sig=remove_manager, tables_data=tables_data, image_paths=image_paths)

            generated_pdf_name = None
        except Exception as e:
            flash(f"Error processing document: {e}", "danger")
            return redirect(url_for("index"))

        user_data = session.get('user', {})
        user_email = user_data.get('email', '')
        is_admin = is_current_user_admin()

        status = 'approved' if is_admin else 'pending'
        if action_type == "reserve":
            status = 'pending'

        if generated_name:
            from services.onedrive import upload_file_to_share, LAST_ERROR
            if os.environ.get("ONEDRIVE_SHARE_URL"):
                success = upload_file_to_share(generated_path, generated_name)
                # We do not strictly need to upload the PDF for previews, but we do if available
                if generated_pdf_name:
                    upload_file_to_share(generated_pdf_path, generated_pdf_name)
                
                # Cleanup preview
                old_preview = session.pop('last_preview_filename', None)
                if old_preview:
                    from services.onedrive import delete_file_from_share
                    delete_file_from_share(old_preview, is_preview=True)
                
                if success:
                    if status == 'approved':
                        flash("Document generated and successfully uploaded to OneDrive!", "success")
                    else:
                        flash("Document generated and uploaded to OneDrive. Pending admin approval.", "info")
                else:
                    import services.onedrive as od
                    if status == 'approved':
                        flash(f"Document generated locally, but OneDrive upload failed. Error: {od.LAST_ERROR}", "danger")
                    else:
                        flash(f"Document pending, but OneDrive upload failed. Error: {od.LAST_ERROR}", "danger")
            else:
                if status == 'approved':
                    flash("Document generated locally. (OneDrive upload not configured.)", "info")
                else:
                    flash("Document uploaded successfully. It is pending admin approval.", "info")

        # Save record to database
        doc = Document(serial_number=serial, full_code=full_code, month=month, year=year, filename=original_name,
                       file_title=file_title, generated_filename=generated_name,
                       generated_pdf_filename=generated_pdf_name, username=user_data.get('name', 'Unknown'),
                       email=user_email, status=status)

        if action_type == "direct_template":
            doc.doc_type = "direct"
            doc.sender = request.form.get("sender", "")
            doc.sender_position = request.form.get("sender_position", "")
            doc.recipient = request.form.get("recipient", "")
            doc.subject = request.form.get("subject", "")
            doc.holder_name = request.form.get("holder_name", "")
            doc.position = request.form.get("position", "")
            doc.is_internal = request.form.get("is_internal") == "1"
            doc.manager_is_sender = request.form.get("manager_is_sender") == "1"
            doc.generated_body = request.form.get("full_body", "")

            # Auto-save contact info
            if doc.recipient:
                existing_rec = Contact.query.filter_by(name=doc.recipient, type="recipient").first()
                if existing_rec:
                    if doc.holder_name:
                        existing_rec.holder_name = doc.holder_name
                    if doc.position:
                        existing_rec.position = doc.position
                else:
                    new_rec = Contact(name=doc.recipient, type="recipient", holder_name=doc.holder_name, position=doc.position)
                    db.session.add(new_rec)
            
            if doc.sender:
                existing_sender = Contact.query.filter_by(name=doc.sender, type="sender").first()
                if not existing_sender:
                    new_sender = Contact(name=doc.sender, type="sender")
                    db.session.add(new_sender)
        elif action_type == "reserve":
            doc.doc_type = "reserved"

        # Store the direct SharePoint Word Online URL
        if generated_name:
            try:
                from services.onedrive import get_word_online_url
                doc.sharepoint_url = get_word_online_url(generated_name)
            except Exception:
                pass

        db.session.add(doc)
        db.session.commit()

        if status == 'approved':
            flash(f"Document number <strong>{full_code}</strong> generated successfully!", "success")
            return redirect(url_for("documents", auto_download=doc.id))
        return redirect(url_for("documents"))

    # Build preview code for the current moment
    now_preview = datetime.now()
    next_serial = get_next_serial(now_preview.month, now_preview.year)
    preview_code = build_full_code(next_serial, now_preview.month, now_preview.year)
    
    contacts = Contact.query.all()
    contacts_data = [
        {
            "name": c.name,
            "type": c.type,
            "holder_name": c.holder_name,
            "position": c.position
        } for c in contacts
    ]
    
    return render_template("index.html", preview_code=preview_code, contacts=contacts_data)


@app.route("/documents")
@login_required
def documents():
    """Documents history page — list all processed documents."""
    user_email = session.get('user', {}).get('email', '')
    is_admin = is_current_user_admin()
    if is_admin:
        all_docs = Document.query.order_by(Document.created_at.desc()).all()
    else:
        all_docs = Document.query.filter_by(email=user_email).order_by(Document.created_at.desc()).all()
    return render_template("documents.html", documents=all_docs, now=datetime.now(), is_admin=is_admin)


@app.route("/download/<int:doc_id>")
@login_required
def download(doc_id):
    """Download the generated (numbered) document."""
    doc = Document.query.get_or_404(doc_id)
    if not doc.generated_filename:
        abort(404)
        
    local_path = os.path.join(GENERATED_FOLDER, doc.generated_filename)
    if not os.path.exists(local_path):
        from services.onedrive import download_file_from_share
        if os.environ.get("ONEDRIVE_SHARE_URL"):
            download_file_from_share(doc.generated_filename, local_path)
            
    if not os.path.exists(local_path):
        abort(404)
        
    return send_from_directory(GENERATED_FOLDER, doc.generated_filename, as_attachment=True,
                               download_name=doc.generated_filename)


@app.route("/download_pdf/<int:doc_id>")
@login_required
def download_pdf(doc_id):
    """Download the generated PDF document."""
    doc = Document.query.get_or_404(doc_id)
    if not doc.generated_pdf_filename:
        flash("No PDF file available for this document.", "warning")
        return redirect(url_for("documents"))
        
    local_path = os.path.join(GENERATED_FOLDER, doc.generated_pdf_filename)
    if not os.path.exists(local_path):
        from services.onedrive import download_file_from_share
        if os.environ.get("ONEDRIVE_SHARE_URL"):
            download_file_from_share(doc.generated_pdf_filename, local_path)
            
    if not os.path.exists(local_path):
        flash("File could not be found locally or on OneDrive.", "danger")
        return redirect(url_for("documents"))
        
    return send_from_directory(GENERATED_FOLDER, doc.generated_pdf_filename, as_attachment=True,
                               download_name=doc.generated_pdf_filename)


@app.route("/view_pdf/<int:doc_id>")
@login_required
def view_pdf(doc_id):
    """View the generated PDF document inline."""
    doc = Document.query.get_or_404(doc_id)
    if not doc.generated_pdf_filename:
        flash("No PDF file available for this document.", "warning")
        return redirect(url_for("documents"))
        
    local_path = os.path.join(GENERATED_FOLDER, doc.generated_pdf_filename)
    if not os.path.exists(local_path):
        from services.onedrive import download_file_from_share
        if os.environ.get("ONEDRIVE_SHARE_URL"):
            download_file_from_share(doc.generated_pdf_filename, local_path)
            
    if not os.path.exists(local_path):
        flash("File could not be found locally or on OneDrive.", "danger")
        return redirect(url_for("documents"))
        
    return send_from_directory(GENERATED_FOLDER, doc.generated_pdf_filename, as_attachment=False)


@app.route("/view_word/<int:doc_id>")
@login_required
def view_word_online(doc_id):
    """Redirect to Word Online to view the document."""
    doc = Document.query.get_or_404(doc_id)
    if not doc.generated_filename:
        flash("No Word file available.", "warning")
        return redirect(url_for("documents"))
        
    import urllib.parse
    filename = urllib.parse.quote(doc.generated_filename)
    url = f"https://horusuni-my.sharepoint.com/personal/aelshafee_horus_edu_eg/_layouts/15/Doc.aspx?sourcedoc=/personal/aelshafee_horus_edu_eg/Documents/ESSIC_Docs/{filename}&action=default"
    return redirect(url)


@app.route("/preview_document", methods=["POST"])
@login_required
def preview_document():
    action_type = request.form.get("action_type")

    import tempfile
    from services.word_editor import replace_placeholder, replace_placeholders
    from services.ai_generator import normalize_recipient
    from services.onedrive import upload_file_to_share, get_word_online_url
    from flask import jsonify

    now = datetime.now()
    serial = get_next_serial(now.month, now.year)
    full_code = build_full_code(serial, now.month, now.year)

    temp_dir = tempfile.mkdtemp()
    temp_docx_path = os.path.join(temp_dir, "preview.docx")

    try:
        if action_type == "upload":
            file = request.files.get("document") or request.files.get("docx_file")
            if not file or file.filename == "":
                return jsonify({"success": False, "error": "No file uploaded"}), 400

            upload_path = os.path.join(temp_dir, "upload.docx")
            file.save(upload_path)
            replace_placeholder(upload_path, temp_docx_path, full_code)

        elif action_type == "direct_template":
            upload_path = os.path.join(BASE_DIR, "template_essic.docx")
            if not os.path.exists(upload_path):
                return jsonify({"success": False, "error": "Template not found"}), 404

            recipient = request.form.get("recipient", "")
            subject = request.form.get("subject", "")
            holder_name = request.form.get("holder_name", "")
            position = request.form.get("position", "")
            full_body = request.form.get("full_body", "")
            sender = request.form.get("sender", "")
            sender_position = request.form.get("sender_position", "")
            is_internal = request.form.get("is_internal") == "1"
            manager_is_sender = request.form.get("manager_is_sender") == "1"

            import re
            def remove_office_code(text):
                if not text: return ""
                text = re.sub(r'\(\s*ESSIC\s*\)', '', text, flags=re.IGNORECASE)
                text = re.sub(r'ESSIC', '', text, flags=re.IGNORECASE)
                return text.strip()

            default_sender = 'مدير مركز الخدمات الإلكترونية والإبداع العلمي (ESSIC)'
            
            # Unconditionally remove ESSIC text from all text fields
            recipient = remove_office_code(recipient)
            subject = remove_office_code(subject)
            holder_name = remove_office_code(holder_name)
            position = remove_office_code(position)
            full_body = remove_office_code(full_body)
            sender = remove_office_code(sender)
            sender_position = remove_office_code(sender_position)
            default_sender = remove_office_code(default_sender)

            sender_value = ""
            sender_name_used = sender if sender else default_sender
            if sender_position:
                sender_top = f"\u202A{sender_name_used} / {sender_position}\u202C"
            else:
                sender_top = f"\u202A{sender_name_used}\u202C"
                
            if is_internal and not manager_is_sender and sender:
                sender_value = f"\u202Bمقدمه لسيادتكم\n{sender}\u202C"

            tables_data_str = request.form.get("tables_data", "[]")
            import json
            try:
                tables_data = json.loads(tables_data_str)
            except:
                tables_data = []

            image_paths = []
            images = request.files.getlist("images")
            if images:
                import uuid
                for img in images:
                    if img.filename:
                        ext = img.filename.split('.')[-1]
                        temp_path = os.path.join(temp_dir, f"temp_{uuid.uuid4().hex}.{ext}")
                        img.save(temp_path)
                        image_paths.append(temp_path)

            replace_placeholders(upload_path, temp_docx_path, {
                "{{CODE_NUMBER}}": full_code,
                "{{SEND_TO}}": ("\u202A" + normalize_recipient(recipient) + "\u202C") if recipient else "",
                "{{SUBJECT}}": "\u202A" + subject + "\u202C",
                "{{STACK_HOLDER}}": "\u202B" + holder_name + "\u202C",
                "{{POSITION}}": "\u202B" + position + "\u202C",
                "{{BODY_TEXT}}": "\u202B" + full_body + "\u202C",
                "{{SENDER}}": sender_value,
                "{{SENDER_TOP}}": sender_top,
                "{{ESSIC}}": "",
            }, remove_manager_sig=(is_internal and not manager_is_sender), tables_data=tables_data, image_paths=image_paths)
        else:
            return jsonify({"success": False, "error": "Invalid action type"}), 400

        # Upload the temporary preview doc to OneDrive
        if os.environ.get("ONEDRIVE_SHARE_URL") or os.environ.get("ONEDRIVE_PREVIEW_SHARE_URL"):
            user_email = session.get('user', {}).get('email', 'unknown')
            username = user_email.split('@')[0]
            timestamp = int(now.timestamp())
            gen_name = f"Preview_{username}_{timestamp}.docx"
            
            success = upload_file_to_share(temp_docx_path, gen_name, is_preview=True)
            if success:
                from services.onedrive import delete_file_from_share
                old_preview = session.get('last_preview_filename')
                if old_preview and old_preview != gen_name:
                    # Attempt to delete the previous preview file to keep OneDrive clean
                    delete_file_from_share(old_preview, is_preview=True)
                session['last_preview_filename'] = gen_name
                
                url = get_word_online_url(gen_name, is_preview=True)
                return jsonify({"success": True, "word_online_url": url})
            else:
                import services.onedrive as od
                return jsonify({"success": False, "error": f"Failed to upload preview to OneDrive. {od.LAST_ERROR}"})
        else:
            return jsonify({"success": False, "error": "OneDrive is not configured. Cannot preview in Word Online."})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/extract_text", methods=["POST"])
@login_required
def extract_text():
    file = request.files.get("docx_file")
    if not file or not allowed_file(file.filename):
        return jsonify({"success": False, "error": "Invalid file"}), 400
        
    try:
        from docx import Document
        doc = Document(file)
        raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        
        from services.ai_generator import extract_document_fields
        user_email = session.get('user', {}).get('email', '')
        
        fields = extract_document_fields(raw_text, user_email)
        
        if fields:
            return jsonify({"success": True, "text": fields.get("full_body", raw_text), "fields": fields})
        else:
            return jsonify({"success": True, "text": raw_text, "fields": None})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/upload_reserved/<int:doc_id>", methods=["POST"])
@login_required
def upload_reserved(doc_id):
    """Upload finalized document for a reserved serial."""
    doc = Document.query.get_or_404(doc_id)
    
    if doc.doc_type != "reserved":
        flash("This document is not eligible for this action.", "danger")
        return redirect(url_for("documents"))
        
    file = request.files.get("docx_file")
    if not file or file.filename == "":
        flash("Please select a Word document to upload.", "warning")
        return redirect(url_for("documents"))
        
    if not allowed_file(file.filename):
        flash("Only .docx files are accepted.", "danger")
        return redirect(url_for("documents"))
        
    try:
        now = datetime.now()
        timestamp = now.strftime("%Y%md_%H%M%S")
        original_name = secure_filename(file.filename)
        generated_name = f"{doc.full_code.replace(' ', '_')}.docx"
        
        generated_path = os.path.join(GENERATED_FOLDER, generated_name)
        file.save(generated_path)
        
        doc.filename = original_name
        doc.generated_filename = generated_name
        doc.generated_pdf_filename = None
        
        is_admin = is_current_user_admin()
        status = 'approved' if is_admin else 'pending'
        
        if status == 'approved':
            from services.onedrive import upload_file_to_share
            if os.environ.get("ONEDRIVE_SHARE_URL"):
                upload_file_to_share(generated_path, generated_name)

        # Store the direct SharePoint Word Online URL
        if generated_name:
            try:
                from services.onedrive import get_word_online_url
                doc.sharepoint_url = get_word_online_url(generated_name)
            except Exception:
                pass

        doc.status = status
        db.session.commit()
        
        if status == 'approved':
            flash("Final document uploaded and approved!", "success")
        else:
            flash("Final document uploaded successfully. It is pending admin approval.", "info")
            
    except Exception as e:
        flash(f"Error processing finalized document: {e}", "danger")
        
    return redirect(url_for("documents"))

@app.route("/edit/<int:doc_id>", methods=["GET", "POST"])
@login_required
def edit(doc_id):
    """Edit a document's serial number and regenerate its full code."""
    doc = Document.query.get_or_404(doc_id)

    user_email = session.get('user', {}).get('email', '')
    is_admin = is_current_user_admin()
    if not is_admin and doc.email != user_email:
        flash("You do not have permission to modify this document.", "danger")
        return redirect(url_for("documents"))

    if request.method == "POST":
        try:
            new_serial = int(request.form["serial_number"])
            if new_serial < 1:
                raise ValueError
        except (ValueError, KeyError):
            flash("Serial number must be a positive integer.", "danger")
            return redirect(url_for("edit", doc_id=doc_id))

        full_body = doc.generated_body
        if doc.doc_type == "direct":
            doc.recipient = request.form.get("recipient", doc.recipient)
            doc.subject = request.form.get("subject", doc.subject)
            doc.holder_name = request.form.get("holder_name", doc.holder_name)
            doc.position = request.form.get("position", doc.position)
            doc.sender = request.form.get("sender", doc.sender)
            doc.is_internal = request.form.get("is_internal") == "1"
            doc.manager_is_sender = request.form.get("manager_is_sender") == "1"
            full_body = request.form.get("full_body", doc.generated_body)
            doc.generated_body = full_body

        # Regenerate full code with new serial
        new_full_code = build_full_code(new_serial, doc.month, doc.year)

        if new_serial != doc.serial_number:
            existing = Document.query.filter_by(month=doc.month, year=doc.year, serial_number=new_serial).first()
            if existing:
                flash(f"Serial number {new_serial} is already used for this month. Please choose another.", "danger")
                return redirect(url_for("edit", doc_id=doc_id))

        # Re-process the Word document with the new code
        upload_path = os.path.join(UPLOAD_FOLDER, doc.filename)
        # If it was generated via direct entry and we don't have the uploaded file, fallback to template
        if doc.doc_type == "direct" and not os.path.exists(upload_path):
            upload_path = os.path.join(BASE_DIR, "template_essic.docx")

        if not os.path.exists(upload_path):
            flash("Original uploaded file not found. Cannot regenerate document.", "danger")
            return redirect(url_for("documents"))

        # Remove old generated file
        if doc.generated_filename:
            old_path = os.path.join(GENERATED_FOLDER, doc.generated_filename)
            if os.path.exists(old_path):
                os.remove(old_path)

        if doc.generated_pdf_filename:
            old_pdf_path = os.path.join(GENERATED_FOLDER, doc.generated_pdf_filename)
            if os.path.exists(old_pdf_path):
                os.remove(old_pdf_path)

        new_generated_name = f"{new_full_code.replace(' ', '_')}.docx"
        new_generated_path = os.path.join(GENERATED_FOLDER, new_generated_name)

        try:
            if doc.doc_type == "direct":
                from services.word_editor import replace_placeholders
                from services.ai_generator import normalize_recipient

                import re
                def remove_office_code(text):
                    if not text: return ""
                    text = re.sub(r'\(\s*ESSIC\s*\)', '', text, flags=re.IGNORECASE)
                    text = re.sub(r'ESSIC', '', text, flags=re.IGNORECASE)
                    return text.strip()

                default_sender = 'مدير مركز الخدمات الإلكترونية والإبداع العلمي (ESSIC)'
                
                # Unconditionally remove ESSIC text from all text fields
                doc.recipient = remove_office_code(doc.recipient)
                doc.subject = remove_office_code(doc.subject)
                doc.holder_name = remove_office_code(doc.holder_name)
                doc.position = remove_office_code(doc.position)
                full_body = remove_office_code(full_body)
                doc.sender = remove_office_code(doc.sender)
                doc.sender_position = remove_office_code(getattr(doc, 'sender_position', ''))
                default_sender = remove_office_code(default_sender)

                body_to_use = full_body
                sender_value = ""
                sender_name_used = doc.sender if doc.sender else default_sender
                if doc.sender_position:
                    sender_top = f"\u202A{sender_name_used} / {doc.sender_position}\u202C"
                else:
                    sender_top = f"\u202A{sender_name_used}\u202C"

                if doc.is_internal and not doc.manager_is_sender and doc.sender:
                    sender_value = f"\u202Bمقدمه لسيادتكم\n{doc.sender}\u202C"

                replace_placeholders(upload_path, new_generated_path, {
                    "{{CODE_NUMBER}}": new_full_code,
                    "{{SEND_TO}}": ("\u202A" + normalize_recipient(doc.recipient) + "\u202C") if doc.recipient else "",
                    "{{SUBJECT}}": "\u202A" + doc.subject + "\u202C",
                    "{{STACK_HOLDER}}": "\u202B" + (doc.holder_name or "") + "\u202C",
                    "{{POSITION}}": "\u202B" + (doc.position or "") + "\u202C",
                    "{{BODY_TEXT}}": "\u202B" + (body_to_use or "") + "\u202C",
                    "{{SENDER}}": sender_value,
                    "{{SENDER_TOP}}": sender_top,
                    "{{ESSIC}}": "",
                }, remove_manager_sig=(doc.is_internal and not doc.manager_is_sender))
            else:
                from services.word_editor import replace_placeholder
                replace_placeholder(upload_path, new_generated_path, new_full_code)

            doc.generated_pdf_filename = None
        except Exception as e:
            flash(f"Error regenerating document: {e}", "danger")
            return redirect(url_for("documents"))

        doc.serial_number = new_serial
        doc.full_code = new_full_code
        doc.generated_filename = new_generated_name
        db.session.commit()

        flash(f"Document updated to <strong>{new_full_code}</strong>.", "success")
        return redirect(url_for("documents"))

    return render_template("edit.html", doc=doc)


@app.route("/delete/<int:doc_id>", methods=["POST"])
@login_required
def delete(doc_id):
    """Delete a document record and its generated file."""
    doc = Document.query.get_or_404(doc_id)

    user_email = session.get('user', {}).get('email', '')
    is_admin = is_current_user_admin()
    if not is_admin and doc.email != user_email:
        flash("You do not have permission to delete this document.", "danger")
        return redirect(url_for("documents"))

    # Remove generated files from disk
    if doc.generated_filename:
        gen_path = os.path.join(GENERATED_FOLDER, doc.generated_filename)
        if os.path.exists(gen_path):
            os.remove(gen_path)

    if doc.generated_pdf_filename:
        gen_pdf_path = os.path.join(GENERATED_FOLDER, doc.generated_pdf_filename)
        if os.path.exists(gen_pdf_path):
            os.remove(gen_pdf_path)

    db.session.delete(doc)
    db.session.commit()

    flash("Document record deleted.", "info")
    return redirect(url_for("documents"))


@app.route("/approve/<int:doc_id>", methods=["POST"])
@login_required
@admin_required
def approve(doc_id):
    """Approve a pending document and upload to OneDrive."""

    doc = Document.query.get_or_404(doc_id)
    if doc.status == 'approved':
        flash("Document is already approved.", "info")
        return redirect(url_for("documents"))

    doc.status = 'approved'
    doc.generated_pdf_filename = None

    # We DO NOT re-upload to OneDrive because it was already uploaded when generated,
    # and the admin might have edited it directly in Word Online.
    flash(f"Document {doc.full_code} approved.", "success")

    # Store the direct SharePoint Word Online URL
    if doc.generated_filename:
        try:
            from services.onedrive import get_word_online_url
            doc.sharepoint_url = get_word_online_url(doc.generated_filename)
        except Exception:
            pass

    db.session.commit()
    return redirect(url_for("documents"))


@app.route("/reject/<int:doc_id>", methods=["POST"])
@login_required
@admin_required
def reject(doc_id):
    """Reject a pending document."""
    doc = Document.query.get_or_404(doc_id)
    if doc.status != 'pending':
        flash("Only pending documents can be rejected.", "warning")
        return redirect(url_for("documents"))

    doc.status = 'rejected'
    db.session.commit()
    flash(f"Document {doc.full_code} has been rejected.", "info")
    return redirect(url_for("documents"))


@app.route("/users")
@login_required
@admin_required
def users():
    """User Management Page."""
    all_users = User.query.order_by(User.created_at.desc()).all()
    return render_template("users.html", users=all_users)


@app.route("/users/role/<email>", methods=["POST"])
@login_required
@admin_required
def toggle_user_role(email):
    user = User.query.filter_by(email=email).first_or_404()
    if user.email in ADMIN_EMAILS:
        flash("Cannot change role of a root admin.", "warning")
    else:
        user.role = 'admin' if user.role == 'user' else 'user'
        db.session.commit()
        flash(f"User {email} is now {user.role}.", "success")
    return redirect(url_for("users"))


@app.route("/users/delete/<email>", methods=["POST"])
@login_required
@admin_required
def delete_user(email):
    user = User.query.filter_by(email=email).first_or_404()
    if user.email in ADMIN_EMAILS:
        flash("Cannot delete a root admin.", "warning")
    else:
        db.session.delete(user)
        db.session.commit()
        flash(f"User {email} has been deleted.", "info")
    return redirect(url_for("users"))


@app.route("/users/add", methods=["POST"])
@login_required
@admin_required
def add_user():
    name = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip().lower()
    role = request.form.get("role", "user").strip()

    if not name or not email:
        flash("Name and email are required.", "warning")
        return redirect(url_for("users"))

    if User.query.filter_by(email=email).first():
        flash("A user with this email already exists.", "warning")
        return redirect(url_for("users"))

    totp_secret = pyotp.random_base32()
    user = User(email=email, name=name, totp_secret=totp_secret, is_setup=False, role=role)
    db.session.add(user)
    db.session.commit()
    flash(f"User {name} added successfully.", "success")
    return redirect(url_for("users"))


@app.route("/users/edit/<email>", methods=["POST"])
@login_required
@admin_required
def edit_user(email):
    user = User.query.filter_by(email=email).first_or_404()

    name = request.form.get("name", "").strip()
    role = request.form.get("role", user.role).strip()
    reset_otp = request.form.get("reset_otp") == "on"

    if name:
        user.name = name

    if user.email not in ADMIN_EMAILS:
        user.role = role
    elif role != 'admin':
        flash("Cannot remove admin role from a root admin.", "warning")

    if reset_otp:
        user.totp_secret = pyotp.random_base32()
        user.is_setup = False
        flash(f"Authenticator reset for {user.name}.", "info")

    db.session.commit()
    flash(f"User {user.email} updated successfully.", "success")
    return redirect(url_for("users"))


# ── AI Content Generator Routes ─────────────────────────────────────────────
#
# Placeholder contract with the .docx/.pages letterhead template:
#   {{CODE_NUMBER}}  -> document code, e.g. "0031 ESSIC 05-2026"
#   {{SEND_TO}}      -> recipient office/department (normalized Arabic)
#   {{SUBJECT}}      -> subject line
#   {{STACK_HOLDER}} -> recipient's formal name/title line, e.g.
#                       "السيد الأستاذ الدكتور / محمد عبدالعال – الموقر"
#   {{POSITION}}     -> recipient's position, e.g. "نائب رئيس جامعة حورس - مصر"
#   {{BODY_TEXT}}    -> generated letter body ONLY (paragraphs, no header/signature)
#
# NOTE: these exact token names (SEND_TO / STACK_HOLDER, in this word order)
# must match whatever .docx/.pages template is actually uploaded. A previous
# version of this code used {{TO_SEND}}/{{HOLDER_STACK}} based on an older
# template file, which didn't match the token names in the template actually
# used in production — those two fields were left unreplaced as literal
# "{{TO_SEND}}"/"{{HOLDER_STACK}}" text. Always grep the real template's XML
# for `{{[A-Z_]*}}` before assuming a placeholder's exact name/order.
#
# The template already renders sender info and the closing signature as
# static content, so those are never re-embedded into {{BODY_TEXT}}.

from services.ai_generator import (generate_arabic_text, build_full_document_text,
                                   get_available_models, normalize_recipient)
from services.word_editor import replace_placeholders


@app.route("/ai/new")
@login_required
def ai_step1():
    return render_template("ai_step1.html")


@app.route("/ai/compose/<doc_type>")
@login_required
def ai_step2(doc_type):
    if doc_type not in ("letter", "request"):
        return redirect(url_for("ai_step1"))
        
    contacts = Contact.query.all()
    contacts_data = [
        {
            "name": c.name,
            "type": c.type,
            "holder_name": c.holder_name,
            "position": c.position
        } for c in contacts
    ]
    
    return render_template("ai_step2.html", doc_type=doc_type, models=get_available_models(), contacts=contacts_data)


@app.route("/ai/generate", methods=["POST"])
@login_required
def ai_generate():
    doc_type = request.form.get("doc_type", "letter")
    sender = request.form.get("sender", "").strip()
    sender_position = request.form.get("sender_position", "").strip()
    recipient = request.form.get("recipient", "").strip()
    subject = request.form.get("subject", "").strip()
    holder_name = request.form.get("holder_name", "").strip()
    position = request.form.get("position", "").strip()
    raw_draft = request.form.get("raw_draft", "").strip()
    model = request.form.get("model")

    import json
    import uuid

    tables_data_str = request.form.get("tables_data", "[]")
    try:
        tables_data = json.loads(tables_data_str)
    except:
        tables_data = []

    temp_image_paths = []
    images = request.files.getlist("images")
    if images:
        for img in images:
            if img.filename:
                ext = img.filename.split('.')[-1]
                temp_path = os.path.join(UPLOAD_FOLDER, f"temp_{uuid.uuid4().hex}.{ext}")
                img.save(temp_path)
                temp_image_paths.append(temp_path)

    if not all([sender, recipient, subject, holder_name, position, raw_draft]):
        flash("يرجى ملء جميع الحقول المطلوبة.", "warning")
        return redirect(url_for("ai_step2", doc_type=doc_type))

    user_email = session['user']['email']
    result = generate_arabic_text(
        doc_type=doc_type, sender=sender, recipient=recipient, subject=subject,
        raw_draft=raw_draft, user_email=user_email,
        num_tables=len(tables_data), num_images=len(temp_image_paths)
    )

    if not result["success"]:
        flash(f"خطأ في توليد المحتوى: {result['error']}", "danger")
        return redirect(url_for("ai_step2", doc_type=doc_type))

    if result.get("warning"):
        flash(result["warning"], "warning")

    full_body = build_full_document_text(generated_body=result["text"])

    return render_template(
        "ai_step3.html", doc_type=doc_type, sender=sender, sender_position=sender_position, 
        recipient=recipient, subject=subject, holder_name=holder_name, position=position, 
        raw_draft=raw_draft, generated_body=result["text"], full_body=full_body, model=model,
        models=get_available_models(), tables_data_json=tables_data_str, 
        image_paths_json=json.dumps(temp_image_paths)
    )


@app.route("/ai/finalize", methods=["POST"])
@login_required
def ai_finalize():
    doc_type = request.form.get("doc_type", "letter")
    sender = request.form.get("sender", "")
    sender_position = request.form.get("sender_position", "")
    recipient = request.form.get("recipient", "")
    subject = request.form.get("subject", "")
    holder_name = request.form.get("holder_name", "")
    position = request.form.get("position", "")
    raw_draft = request.form.get("raw_draft", "")
    full_body = request.form.get("full_body", "")
    model = request.form.get("model")

    tables_data_str = request.form.get("tables_data", "[]")
    import json
    try:
        tables_data = json.loads(tables_data_str)
    except:
        tables_data = []

    image_paths_str = request.form.get("image_paths", "[]")
    try:
        image_paths = json.loads(image_paths_str)
    except:
        image_paths = []

    upload_path = os.path.join(BASE_DIR, "template_essic.docx")
    if not os.path.exists(upload_path):
        flash("قالب النظام الأساسي غير موجود.", "danger")
        return redirect(url_for("ai_step1"))
        
    original_name = "template_essic.docx"

    now = datetime.utcnow()
    serial = get_next_serial(now.month, now.year)
    code = build_full_code(serial, now.month, now.year)

    stem = "template_essic"
    gen_name = f"{code.replace(' ', '_')}.docx"
    gen_path = os.path.join(GENERATED_FOLDER, gen_name)

    import re
    def remove_office_code(text):
        if not text: return ""
        text = re.sub(r'\(\s*ESSIC\s*\)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'ESSIC', '', text, flags=re.IGNORECASE)
        return text.strip()

    default_sender = 'مدير مركز الخدمات الإلكترونية والإبداع العلمي (ESSIC)'
    
    # Unconditionally remove ESSIC text from all text fields
    recipient = remove_office_code(recipient)
    subject = remove_office_code(subject)
    holder_name = remove_office_code(holder_name)
    position = remove_office_code(position)
    full_body = remove_office_code(full_body)
    sender = remove_office_code(sender)
    sender_position = remove_office_code(sender_position)
    sender_name_used = sender if sender else default_sender
    if sender_position:
        sender_top = f"\u202A{sender_name_used} / {sender_position}\u202C"
    else:
        sender_top = f"\u202A{sender_name_used}\u202C"

    try:
        replace_placeholders(upload_path, gen_path, {
            "{{CODE_NUMBER}}": code,
            "{{SEND_TO}}": "\u202A" + normalize_recipient(recipient) + "\u202C" if recipient else "",
            "{{SUBJECT}}": "\u202A" + subject + "\u202C",
            "{{STACK_HOLDER}}": "\u202B" + holder_name + "\u202C",
            "{{POSITION}}": "\u202B" + position + "\u202C",
            "{{BODY_TEXT}}": "\u202B" + full_body + "\u202C",
            "{{SENDER}}": "", # Manager signature kept
            "{{SENDER_TOP}}": sender_top,
            "{{ESSIC}}": "",
        }, tables_data=tables_data, image_paths=image_paths)

        gen_pdf_name = None

    except Exception as e:
        flash(f"خطأ في معالجة المستند: {e}", "danger")
        return redirect(url_for("ai_step1"))

    user_email = session['user']['email']
    user_name = session['user']['name']
    status = 'approved' if is_current_user_admin() else 'pending'

    doc = Document(serial_number=serial, full_code=code, month=now.month, year=now.year, doc_type=doc_type,
                   sender=sender, sender_position=sender_position, recipient=recipient, subject=subject, holder_name=holder_name, position=position,
                   raw_draft=raw_draft, generated_body=full_body,
                   filename=original_name, generated_filename=gen_name, generated_pdf_filename=gen_pdf_name,
                   ai_model=model, username=user_name, email=user_email,
                   status=status, )
    db.session.add(doc)
    db.session.commit()

    if status == 'approved':
        from services.onedrive import upload_file_to_onedrive
        try:
            upload_file_to_onedrive(gen_path, gen_name)
        except Exception as e:
            flash(f"Error uploading to OneDrive: {e}", "danger")

    flash(f"تم إنشاء المستند بنجاح — الرمز: <strong>Code No {code}</strong>", "success")
    return redirect(url_for("documents"))


@app.route("/api/ai/regenerate", methods=["POST"])
@login_required
def api_ai_regenerate():
    data = request.get_json()
    user_email = session['user']['email']
    result = generate_arabic_text(
        doc_type=data.get("doc_type", "letter"), sender=data.get("sender", ""),
        recipient=data.get("recipient", ""), subject=data.get("subject", ""),
        raw_draft=data.get("raw_draft", ""), user_email=user_email,
        num_tables=data.get("num_tables", 0), num_images=data.get("num_images", 0)
    )
    if result["success"]:
        full_body = build_full_document_text(generated_body=result["text"])
        return jsonify({"success": True, "body": result["text"], "full_body": full_body,
                        "warning": result.get("warning")})
    return jsonify({"success": False, "error": result["error"]})

@app.route("/api/ai/format", methods=["POST"])
@login_required
def api_ai_format_document():
    data = request.get_json()
    user_email = session['user']['email']
    full_body = data.get("full_body", "")
    has_table = data.get("has_table", False)
    has_images = data.get("has_images", False)
    
    from services.ai_generator import format_document_with_ai
    result = format_document_with_ai(full_body, has_table, has_images, user_email)
    
    if result.get("success"):
        return jsonify({"success": True, "full_body": result["text"]})
    return jsonify({"success": False, "error": result.get("error", "Unknown error")})


# ─── Entry Point ─────────────────────────────────────────────────────────────

@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings_page():
    user_email = session['user']['email']
    user = User.query.filter_by(email=user_email).first()

    if not user:
        flash("User not found in database.", "danger")
        return redirect(url_for('documents'))

    from services.ai_generator import encrypt_api_key, get_user_keys

    if request.method == "POST":
        gemini = request.form.get("gemini_key", "").strip()
        groq = request.form.get("groq_key", "").strip()
        deepseek = request.form.get("deepseek_key", "").strip()

        user.gemini_api_key = encrypt_api_key(gemini) if gemini else ""
        user.groq_api_key = encrypt_api_key(groq) if groq else ""
        user.deepseek_api_key = encrypt_api_key(deepseek) if deepseek else ""

        db.session.commit()
        flash("API Keys saved securely.", "success")
        return redirect(url_for('settings_page'))

    user_keys = get_user_keys(user_email)
    return render_template("settings.html", user_keys=user_keys)


# ─── Contacts Management ─────────────────────────────────────────────────────

@app.route("/contacts")
@login_required
def contacts():
    all_contacts = Contact.query.order_by(Contact.name).all()
    return render_template("contacts.html", contacts=all_contacts)


@app.route("/contacts/add", methods=["POST"])
@login_required
@admin_required
def add_contact():
    name = request.form.get("name", "").strip()
    c_type = request.form.get("type", "").strip()
    holder_name = request.form.get("holder_name", "").strip()
    position = request.form.get("position", "").strip()
    
    if name:
        contact = Contact(
            name=name,
            type=c_type if c_type else None,
            holder_name=holder_name,
            position=position
        )
        db.session.add(contact)
        db.session.commit()
        flash("Contact added successfully.", "success")
    else:
        flash("Name is required.", "danger")
        
    return redirect(url_for('contacts'))


@app.route("/contacts/edit/<int:contact_id>", methods=["POST"])
@login_required
@admin_required
def edit_contact(contact_id):
    contact = Contact.query.get_or_404(contact_id)
    contact.name = request.form.get("name", "").strip()
    c_type = request.form.get("type", "").strip()
    contact.type = c_type if c_type else None
    contact.holder_name = request.form.get("holder_name", "").strip()
    contact.position = request.form.get("position", "").strip()
    
    db.session.commit()
    flash("Contact updated successfully.", "success")
    return redirect(url_for('contacts'))


@app.route("/contacts/delete/<int:contact_id>", methods=["POST"])
@login_required
@admin_required
def delete_contact(contact_id):
    contact = Contact.query.get_or_404(contact_id)
    db.session.delete(contact)
    db.session.commit()
    flash("Contact deleted successfully.", "success")
    return redirect(url_for('contacts'))
@app.route("/uploads/<path:filename>")
@login_required
def serve_upload(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

if __name__ == "__main__":
    app.run(debug=true, host="0.0.0.0", port=5010)