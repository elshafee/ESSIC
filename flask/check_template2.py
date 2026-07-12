from docx import Document
doc = Document("template_essic.docx")
# Check if there are paragraphs after the manager name
for i, p in enumerate(doc.paragraphs):
    print(f"P{i}: text=[{p.text}]")
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# Check if {{SENDER}} is in the template
has_sender = False
for p in doc.paragraphs:
    if "{{SENDER}}" in p.text:
        has_sender = True
        print(f"\n{{SENDER}} found in paragraph: [{p.text}]")
if not has_sender:
    print("\n{{SENDER}} NOT found in any paragraph!")
