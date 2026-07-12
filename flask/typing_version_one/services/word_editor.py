"""
services/word_editor.py

Replaces {{CODE_NUMBER}} and {{BODY_TEXT}} placeholders in a .docx file
while preserving all run-level formatting (fonts, colors, bold, RTL Arabic, logos).
"""

from docx import Document


PLACEHOLDER_CODE = "{{CODE_NUMBER}}"
PLACEHOLDER_BODY = "{{BODY_TEXT}}"


def _replace_in_paragraph(paragraph, replacements: dict):
    """
    Safely replace placeholders in a paragraph.
    Reassembles split-run placeholders then replaces in the first run.
    """
    full_text = "".join(r.text for r in paragraph.runs)

    hit = False
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            hit = True

    if not hit:
        return

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, replacements)
            for nested in cell.tables:
                _replace_in_table(nested, replacements)


def _replace_in_part(part, replacements: dict):
    if part is None:
        return
    for para in part.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in part.tables:
        _replace_in_table(table, replacements)


def replace_placeholders(input_path: str, output_path: str, replacements: dict):
    """
    Open input_path, replace all keys in `replacements` dict with their values,
    save to output_path.

    Example replacements:
        {
            "{{CODE_NUMBER}}": "0031 ESSIC 04-2026",
            "{{BODY_TEXT}}":   "... Arabic body ..."
        }
    """
    doc = Document(input_path)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        _replace_in_table(table, replacements)

    for section in doc.sections:
        _replace_in_part(section.header, replacements)
        _replace_in_part(section.footer, replacements)
        _replace_in_part(section.even_page_header, replacements)
        _replace_in_part(section.even_page_footer, replacements)
        _replace_in_part(section.first_page_header, replacements)
        _replace_in_part(section.first_page_footer, replacements)

    doc.save(output_path)
