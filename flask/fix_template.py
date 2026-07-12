from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import copy

doc = Document("template_essic.docx")

# Find P10 (BODY_TEXT) and check spacing
for i, p in enumerate(doc.paragraphs):
    print(f"P{i}: text=[{p.text[:60] if p.text.strip() else 'EMPTY'}]")
    if p.paragraph_format.space_after:
        print(f"  space_after: {p.paragraph_format.space_after}")
    if p.paragraph_format.space_before:
        print(f"  space_before: {p.paragraph_format.space_before}")

# Remove excessive empty paragraphs between CODE_NUMBER and عناية
# P3, P4, P5, P6 are all empty - remove P4, P5, P6 (keep one empty line P3)
# Actually, let's just remove extra spacing after BODY_TEXT 

# Set body text paragraph space after to 0
for p in doc.paragraphs:
    if "{{BODY_TEXT}}" in p.text:
        p.paragraph_format.space_after = Pt(6)
        print(f"\nSet BODY_TEXT space_after to 6pt")

# Now add {{SENDER}} paragraph AFTER the manager signature block (P14)
# This will be for employee signature when internal + not manager
last_para = doc.paragraphs[-1]  # P14 - manager name
print(f"\nLast paragraph: [{last_para.text[:60]}]")

# Add a new paragraph after the last one for {{SENDER}}
new_para = doc.add_paragraph()
run = new_para.add_run("{{SENDER}}")
# Copy font from the manager line
run.font.name = 'Arial Unicode MS'
run.font.size = Pt(19)
run.font.bold = True
rPr = run._element.get_or_add_rPr()
rFonts = rPr.get_or_add_rFonts()
rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
rFonts.set(qn('w:cs'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')

# Set RTL for the new paragraph
from docx.oxml import OxmlElement
pPr = new_para._element.get_or_add_pPr()
bidi = OxmlElement('w:bidi')
pPr.append(bidi)
jc = OxmlElement('w:jc')
jc.set(qn('w:val'), 'right')
pPr.append(jc)

doc.save("template_essic.docx")
print("\nTemplate updated with {{SENDER}} paragraph.")
