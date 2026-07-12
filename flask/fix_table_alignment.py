from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("template_essic.docx")

# The value cells are column 0 (left column in RTL layout = the actual value side)
# Let's print and fix all table cells
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                text = p.text.strip()
                # These are the value cells (column 0) containing placeholders
                if any(ph in text for ph in ["{{SENDER_TOP}}", "{{SEND_TO}}", "{{SUBJECT}}"]):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"Set RIGHT align on T{t_idx}R{r_idx}C{c_idx}: {text[:30]}")

doc.save("template_essic.docx")
print("Saved.")
