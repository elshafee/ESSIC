from docx import Document
doc = Document("template_essic.docx")
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if "{{SENDER}}" in p.text:
                    p.text = p.text.replace("{{SENDER}}", "{{SENDER_TOP}}")
                    # restore RTL font correctly
                    for run in p.runs:
                        run.font.name = 'Arial Unicode MS'
doc.save("template_essic.docx")
print("Template updated with SENDER_TOP")
