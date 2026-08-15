"""
services/word_editor.py

Safe placeholder replacement in Word (.docx) documents.
Replaces {{CODE_NUMBER}} in:
  - paragraphs
  - table cells
  - headers
  - footers

Preserves all formatting, fonts, Arabic text, logos, and layout.
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PLACEHOLDER = "{{CODE_NUMBER}}"


def _replace_in_paragraph(paragraph, replacement: str):
    """
    Replace placeholder in a paragraph while preserving run-level formatting.

    Word sometimes splits a placeholder across multiple runs (e.g. {{CODE_ and NUMBER}}).
    We reassemble the full paragraph text, detect the placeholder, then rebuild runs
    carefully to avoid formatting loss.
    """
    full_text = "".join(run.text for run in paragraph.runs)

    if PLACEHOLDER not in full_text:
        return  # nothing to do

    # Replace in the combined text
    new_text = full_text.replace(PLACEHOLDER, replacement)

    # Strategy: put all new text into the FIRST run, clear the rest.
    # This preserves the first run's character formatting for the replaced text.
    if not paragraph.runs:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, replacement: str):
    """Iterate every cell and every paragraph inside a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacement)
            # Tables can be nested
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacement)


def _replace_in_header_footer(section_part, replacement: str):
    """Replace placeholder inside a header or footer section part."""
    if section_part is None:
        return
    for paragraph in section_part.paragraphs:
        _replace_in_paragraph(paragraph, replacement)
    for table in section_part.tables:
        _replace_in_table(table, replacement)


def replace_placeholder(input_path: str, output_path: str, replacement: str):
    """
    Opens a .docx file, replaces all occurrences of {{CODE_NUMBER}} with
    the provided replacement string, and saves to output_path.

    Args:
        input_path:   Path to the source .docx template.
        output_path:  Path where the modified document will be saved.
        replacement:  The code string to substitute (e.g. '0031 ESSIC 05-2026').
    """
    doc = Document(input_path)

    # 1. Body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacement)

    # 2. Body tables
    for table in doc.tables:
        _replace_in_table(table, replacement)

    # 3. Headers and footers across all sections
    for section in doc.sections:
        _replace_in_header_footer(section.header, replacement)
        _replace_in_header_footer(section.footer, replacement)
        _replace_in_header_footer(section.even_page_header, replacement)
        _replace_in_header_footer(section.even_page_footer, replacement)
        _replace_in_header_footer(section.first_page_footer, replacement)

    doc.save(output_path)


def replace_placeholders(input_path: str, output_path: str, replacements: dict, remove_manager_sig: bool = False, tables_data: list = None, image_paths: list = None):
    """
    Opens a .docx file, replaces multiple occurrences of keys in `replacements` dict
    with their corresponding values, and saves to output_path.
    If remove_manager_sig is True, it removes the hardcoded manager signature lines.
    """
    doc = Document(input_path)

    def _replace_in_p(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        modified = False
        remove_p = False
        modified_body = False

        for k, v in replacements.items():
            if k in full_text:
                if k == "{{BODY_TEXT}}":
                    modified_body = True
                    # Clean up Body Text
                    raw_v = v
                    prefix = ""
                    suffix = ""
                    if raw_v.startswith("\u202B") or raw_v.startswith("\u202A"):
                        prefix = raw_v[0]
                        raw_v = raw_v[1:]
                    if raw_v.endswith("\u202C"):
                        suffix = raw_v[-1]
                        raw_v = raw_v[:-1]
                        
                    # Enforce limits (140 words, 900 chars)
                    MAX_WORDS = 141
                    MAX_CHARS = 900
                    raw_v = raw_v.strip()
                    words = raw_v.split()
                    if len(words) > MAX_WORDS:
                        raw_v = " ".join(words[:MAX_WORDS])
                    if len(raw_v) > MAX_CHARS:
                        raw_v = raw_v[:MAX_CHARS].rsplit(" ", 1)[0]
                        
                    # Filter out empty lines
                    lines = [line.strip() for line in raw_v.split("\n") if line.strip()]
                    raw_v = "\n".join(lines)
                    
                    v = prefix + raw_v + suffix

                full_text = full_text.replace(k, v)
                modified = True

        if modified:
            if not full_text.strip():
                paragraph._element.getparent().remove(paragraph._element)
            elif paragraph.runs:
                if modified_body:
                    from html.parser import HTMLParser
                    import re
                    
                    class DocxHTMLParser(HTMLParser):
                        def __init__(self):
                            super().__init__()
                            self.paragraphs = []
                            self.current_paragraph = []
                            self.current_format = {'bold': False, 'italic': False, 'underline': False}
                            
                        def handle_starttag(self, tag, attrs):
                            if tag in ['strong', 'b']:
                                self.current_format['bold'] = True
                            elif tag in ['em', 'i']:
                                self.current_format['italic'] = True
                            elif tag in ['u']:
                                self.current_format['underline'] = True
                            elif tag in ['p', 'li', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                if self.current_paragraph:
                                    self.paragraphs.append(self.current_paragraph)
                                    self.current_paragraph = []
                                if tag == 'li':
                                    self.current_paragraph.append({
                                        'text': '• ',
                                        'bold': False,
                                        'italic': False,
                                        'underline': False
                                    })
                                
                        def handle_endtag(self, tag):
                            if tag in ['strong', 'b']:
                                self.current_format['bold'] = False
                            elif tag in ['em', 'i']:
                                self.current_format['italic'] = False
                            elif tag in ['u']:
                                self.current_format['underline'] = False
                            elif tag in ['p', 'li', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                if self.current_paragraph:
                                    self.paragraphs.append(self.current_paragraph)
                                    self.current_paragraph = []
                            elif tag == 'br':
                                if self.current_paragraph:
                                    self.paragraphs.append(self.current_paragraph)
                                    self.current_paragraph = []

                        def handle_data(self, data):
                            text = data.replace('\n', ' ').replace('\r', '')
                            if not text.strip() and len(self.current_paragraph) == 0:
                                return
                            self.current_paragraph.append({
                                'text': text,
                                'bold': self.current_format['bold'],
                                'italic': self.current_format['italic'],
                                'underline': self.current_format['underline']
                            })

                    if '<p' not in full_text and '<br' not in full_text and '<div' not in full_text:
                        full_text = full_text.replace('\n', '<br>')
                        
                    parser = DocxHTMLParser()
                    parser.feed(full_text)
                    if parser.current_paragraph:
                        parser.paragraphs.append(parser.current_paragraph)
                    
                    parsed_paragraphs = parser.paragraphs
                    
                    # Clear original paragraph text completely
                    for run in paragraph.runs:
                        run.text = ""
                        
                    # Save base run properties from the very first run, if any
                    base_rPr = None
                    if paragraph.runs and paragraph.runs[0]._element.rPr is not None:
                        base_rPr = copy.deepcopy(paragraph.runs[0]._element.rPr)
                        
                    current_p = paragraph
                    
                    for p_data in parsed_paragraphs:
                        if current_p.runs and any(r.text for r in current_p.runs):
                            new_p_element = OxmlElement('w:p')
                            current_p._element.addnext(new_p_element)
                            current_p = Paragraph(new_p_element, current_p._parent)
                            
                        current_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        fmt = current_p.paragraph_format
                        fmt.space_before = Pt(0)
                        fmt.space_after = Pt(0)
                        fmt.line_spacing = 1.0

                        pPr = current_p._element.get_or_add_pPr()
                        bidi = OxmlElement('w:bidi')
                        bidi.set(qn('w:val'), '1')
                        pPr.append(bidi)
                        
                        for run_info in p_data:
                            text_part = run_info['text']
                            pieces = re.split(r"(\{\{TABLE_\d+\}\}|\{\{IMAGE_\d+\}\})", text_part)
                            
                            for part in pieces:
                                if not part:
                                    continue
                                    
                                table_match = re.fullmatch(r"\{\{TABLE_(\d+)\}\}", part)
                                if table_match:
                                    if tables_data:
                                        try:
                                            idx = int(table_match.group(1)) - 1
                                            if 0 <= idx < len(tables_data):
                                                t_data = tables_data[idx]
                                                if t_data:
                                                    from docx.enum.table import WD_TABLE_ALIGNMENT
                                                    new_table = doc.add_table(rows=len(t_data), cols=len(t_data[0]))
                                                    
                                                    new_table.autofit = True
                                                    
                                                    try:
                                                        new_table.style = 'Table Grid'
                                                    except:
                                                        pass

                                                    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                                    tblPr = new_table._element.xpath('w:tblPr')
                                                    if tblPr:
                                                        tblW = OxmlElement('w:tblW')
                                                        tblW.set(qn('w:w'), '5000')
                                                        tblW.set(qn('w:type'), 'pct')
                                                        tblPr[0].append(tblW)
                                                        
                                                        bidiVisual = OxmlElement('w:bidiVisual')
                                                        tblPr[0].append(bidiVisual)
                                                        
                                                        tblBorders = OxmlElement('w:tblBorders')
                                                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                                            border = OxmlElement(f'w:{border_name}')
                                                            border.set(qn('w:val'), 'single')
                                                            border.set(qn('w:sz'), '4')
                                                            border.set(qn('w:space'), '0')
                                                            border.set(qn('w:color'), '343176')
                                                            tblBorders.append(border)
                                                        tblPr[0].append(tblBorders)

                                                    for r_idx, row_data in enumerate(t_data):
                                                        for c_idx, cell_value in enumerate(row_data):
                                                            cell = new_table.cell(r_idx, c_idx)
                                                            cell.text = ""
                                                            p = cell.paragraphs[0]
                                                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                                            pPr = p._element.get_or_add_pPr()
                                                            bidi = OxmlElement('w:bidi')
                                                            bidi.set(qn('w:val'), '1')
                                                            pPr.append(bidi)
                                                            
                                                            run = p.add_run(str(cell_value))
                                                            
                                                            rPr = run._element.get_or_add_rPr()
                                                            rtl = OxmlElement('w:rtl')
                                                            rtl.set(qn('w:val'), '1')
                                                            rPr.append(rtl)
                                                            
                                                            tcPr = cell._element.get_or_add_tcPr()
                                                            shd = OxmlElement('w:shd')
                                                            shd.set(qn('w:val'), 'clear')
                                                            shd.set(qn('w:color'), 'auto')
                                                            
                                                            if r_idx == 0:
                                                                shd.set(qn('w:fill'), '343176')
                                                                run.bold = True
                                                                from docx.shared import RGBColor
                                                                run.font.color.rgb = RGBColor(255, 255, 255)
                                                            else:
                                                                shd.set(qn('w:fill'), 'F4F6F9')
                                                                
                                                            tcPr.append(shd)
                                                    tbl_xml = new_table._element
                                                    tbl_xml.getparent().remove(tbl_xml)
                                                    
                                                    dummy_before_xml = OxmlElement('w:p')
                                                    current_p._element.addnext(dummy_before_xml)
                                                    dummy_before_xml.addnext(tbl_xml)
                                                    
                                                    dummy_p_xml = OxmlElement('w:p')
                                                    tbl_xml.addnext(dummy_p_xml)
                                                    current_p = Paragraph(dummy_p_xml, current_p._parent)
                                                    
                                                    # Re-apply RTL and spacing to the new paragraph
                                                    current_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                                    fmt = current_p.paragraph_format
                                                    fmt.space_before = Pt(0)
                                                    fmt.space_after = Pt(0)
                                                    fmt.line_spacing = 1.0
                                                    pPr = current_p._element.get_or_add_pPr()
                                                    bidi = OxmlElement('w:bidi')
                                                    bidi.set(qn('w:val'), '1')
                                                    pPr.append(bidi)
                                        except Exception as e:
                                            print("Table error:", e)
                                    continue
                                    
                                img_match = re.fullmatch(r"\{\{IMAGE_(\d+)\}\}", part)
                                if img_match:
                                    if image_paths:
                                        try:
                                            idx = int(img_match.group(1)) - 1
                                            if 0 <= idx < len(image_paths):
                                                img_path = image_paths[idx]
                                                from docx.shared import Inches
                                                img_p_xml = OxmlElement('w:p')
                                                current_p._element.addnext(img_p_xml)
                                                img_p = Paragraph(img_p_xml, current_p._parent)
                                                img_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                                run = img_p.add_run()
                                                run.add_picture(img_path, width=Inches(3.0))
                                                
                                                dummy_p_xml = OxmlElement('w:p')
                                                img_p_xml.addnext(dummy_p_xml)
                                                current_p = Paragraph(dummy_p_xml, current_p._parent)
                                                
                                                # Re-apply RTL and spacing to the new paragraph
                                                current_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                                fmt = current_p.paragraph_format
                                                fmt.space_before = Pt(0)
                                                fmt.space_after = Pt(0)
                                                fmt.line_spacing = 1.0
                                                pPr = current_p._element.get_or_add_pPr()
                                                bidi = OxmlElement('w:bidi')
                                                bidi.set(qn('w:val'), '1')
                                                pPr.append(bidi)
                                        except Exception as e:
                                            print("Image error:", e)
                                    continue
                                    
                                run = current_p.add_run(part)
                                if base_rPr is not None:
                                    run._element.append(copy.deepcopy(base_rPr))
                                    
                                rPr = run._element.get_or_add_rPr()
                                rtl = OxmlElement('w:rtl')
                                rtl.set(qn('w:val'), '1')
                                rPr.append(rtl)
                                    
                                if run_info['bold']:
                                    run.bold = True
                                if run_info['italic']:
                                    run.italic = True
                                if run_info['underline']:
                                    run.underline = True
                    
                else:
                    # Normal replacement without paragraph splitting
                    saved_alignment = paragraph.alignment
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    if saved_alignment is not None:
                        paragraph.alignment = saved_alignment

    def _replace_in_t(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_p(paragraph)
                for nested_table in cell.tables:
                    _replace_in_t(nested_table)

    def _replace_in_hf(section_part):
        if section_part is None:
            return
        for paragraph in section_part.paragraphs:
            _replace_in_p(paragraph)
        for table in section_part.tables:
            _replace_in_t(table)

    # If employee is sender, remove manager signature paragraphs
    if remove_manager_sig:
        manager_texts = [
            "مدير مركز الخدمات الإلكترونية",
            "ومدير مركز الخدمات الإلكترونية",
            "محمد كمال عبد السلام"
        ]

        def _remove_manager_p(p):
            text = p.text
            if any(m_text in text for m_text in manager_texts):
                p._element.getparent().remove(p._element)

        for p in doc.paragraphs:
            _remove_manager_p(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _remove_manager_p(p)
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for nested_p in nested_cell.paragraphs:
                                    _remove_manager_p(nested_p)

    for paragraph in doc.paragraphs:
        _replace_in_p(paragraph)
    for table in doc.tables:
        _replace_in_t(table)
    for section in doc.sections:
        _replace_in_hf(section.header)
        _replace_in_hf(section.footer)
        _replace_in_hf(section.even_page_header)
        _replace_in_hf(section.even_page_footer)
        _replace_in_hf(section.first_page_header)
        _replace_in_hf(section.first_page_footer)

    doc.save(output_path)