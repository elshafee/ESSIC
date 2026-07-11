from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document("template_essic.docx")

def apply_font(run):
    run.font.name = 'Arial Unicode MS'
    run.font.size = Pt(19)
    run.font.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
    rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')

for p in doc.paragraphs:
    if "{{BODY_TEXT}}" in p.text or "{{SENDER}}" in p.text:
        for r in p.runs:
            apply_font(r)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "{{BODY_TEXT}}" in p.text or "{{SENDER}}" in p.text:
                    for r in p.runs:
                        apply_font(r)

doc.save("template_essic.docx")
print("Template font updated.")
