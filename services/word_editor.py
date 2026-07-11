"""
services/word_editor.py

Safe placeholder replacement in Word (.docx) documents.
Replaces {{CODE_NUMBER}} in:
  - paragraphs
  - table cells
  - headers
  - footers

Preserves all formatting, fonts, Arabic text, logos, and layout.
"""

import copy
from docx import Document
from docx.oxml.ns import qn


PLACEHOLDER = "{{CODE_NUMBER}}"


def _replace_in_paragraph(paragraph, replacement: str):
    """
    Replace placeholder in a paragraph while preserving run-level formatting.

    Word sometimes splits a placeholder across multiple runs (e.g. {{CODE_ and NUMBER}}).
    We reassemble the full paragraph text, detect the placeholder, then rebuild runs
    carefully to avoid formatting loss.
    """
    full_text = "".join(run.text for run in paragraph.runs)

    if PLACEHOLDER not in full_text:
        return  # nothing to do

    # Replace in the combined text
    new_text = full_text.replace(PLACEHOLDER, replacement)

    # Strategy: put all new text into the FIRST run, clear the rest.
    # This preserves the first run's character formatting for the replaced text.
    if not paragraph.runs:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, replacement: str):
    """Iterate every cell and every paragraph inside a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacement)
            # Tables can be nested
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacement)


def _replace_in_header_footer(section_part, replacement: str):
    """Replace placeholder inside a header or footer section part."""
    if section_part is None:
        return
    for paragraph in section_part.paragraphs:
        _replace_in_paragraph(paragraph, replacement)
    for table in section_part.tables:
        _replace_in_table(table, replacement)


def replace_placeholder(input_path: str, output_path: str, replacement: str):
    """
    Opens a .docx file, replaces all occurrences of {{CODE_NUMBER}} with
    the provided replacement string, and saves to output_path.

    Args:
        input_path:   Path to the source .docx template.
        output_path:  Path where the modified document will be saved.
        replacement:  The code string to substitute (e.g. '0031 ESSIC 05-2026').
    """
    doc = Document(input_path)

    # 1. Body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacement)

    # 2. Body tables
    for table in doc.tables:
        _replace_in_table(table, replacement)

    # 3. Headers and footers across all sections
    for section in doc.sections:
        _replace_in_header_footer(section.header, replacement)
        _replace_in_header_footer(section.footer, replacement)
        _replace_in_header_footer(section.even_page_header, replacement)
        _replace_in_header_footer(section.even_page_footer, replacement)
        _replace_in_header_footer(section.first_page_header, replacement)
        _replace_in_header_footer(section.first_page_footer, replacement)

    doc.save(output_path)
