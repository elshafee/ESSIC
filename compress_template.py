from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

doc = Document("template_essic.docx")

# Clean up excessive empty paragraphs at the top.
# P0-P3 and P5-P10 seem to have many empty ones.
# Let's remove multiple consecutive empty paragraphs.
empty_count = 0
for p in list(doc.paragraphs):
    if not p.text.strip() and "{{BODY_TEXT}}" not in p.text and "{{CODE_NUMBER}}" not in p.text:
        empty_count += 1
        if empty_count > 1: # Keep at most 1 empty line
            p._element.getparent().remove(p._element)
    else:
        empty_count = 0

# Now apply minimum spacing to all remaining paragraphs
for p in doc.paragraphs:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0

doc.save("template_essic.docx")
print("Template compressed successfully.")
