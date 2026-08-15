import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), "flask"))

from services.word_editor import replace_placeholders
from docx import Document

input_docx = "flask/template_essic.docx"
output_docx = "flask/test_output.docx"

replace_placeholders(input_docx, output_docx, {
    "{{ESSIC}}": "REPLACED_ESSIC",
    "{{SENDER_TOP}}": "TEST_SENDER",
})

doc = Document(output_docx)
for p in doc.sections[0].header.paragraphs:
    print("HF P:", p.text)
for t in doc.sections[0].header.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                print("HF T P:", p.text)
