import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), "flask"))
from docx import Document

input_docx = "flask/test_output.docx"
doc = Document(input_docx)

print("--- BODY PARAGRAPHS ---")
for p in doc.paragraphs:
    if "ESSIC" in p.text: print("BODY P:", p.text)

print("--- BODY TABLES ---")
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if "ESSIC" in p.text: print("BODY TABLE P:", p.text)

print("--- HEADERS ---")
for section in doc.sections:
    for hdr in [section.header, section.first_page_header]:
        if not hdr: continue
        for p in hdr.paragraphs:
            if "ESSIC" in p.text: print("HEADER P:", p.text)
        for t in hdr.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        if "ESSIC" in p.text: print("HEADER TABLE P:", p.text)
